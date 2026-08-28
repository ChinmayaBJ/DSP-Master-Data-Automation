"""Convert PROGRESS_REPORT.md -> .docx (self-contained markdown subset parser)."""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HERE = Path(__file__).parent
SRC = str(_HERE / "PROGRESS_REPORT.md")
OUT = str(_HERE / "PROGRESS_REPORT.docx")

doc = Document()
# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

MONO = "Consolas"
CODE_GRAY = RGBColor(0x33, 0x33, 0x33)

def add_runs(paragraph, text):
    """Parse inline **bold**, `code`, [txt](url); add as runs."""
    # split keeping delimiters via a token regex
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = MONO; r.font.color.rgb = CODE_GRAY
        else:  # link [txt](url) -> keep visible text only
            txt = re.match(r"\[([^\]]+)\]", tok).group(1)
            r = paragraph.add_run(txt); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO; run.font.size = Pt(9.5); run.font.color.rgb = CODE_GRAY

def add_table(rows):
    # rows: list of list-of-cells (strings). first row is header.
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            cell_txt = row[ci] if ci < len(row) else ""
            para = cells[ci].paragraphs[0]
            add_runs(para, cell_txt.strip())
            if ri == 0:
                for run in para.runs:
                    run.bold = True

def split_table_row(line):
    line = line.strip()
    if line.startswith("|"): line = line[1:]
    if line.endswith("|"): line = line[:-1]
    return [c.strip() for c in line.split("|")]

with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # fenced code block
    if stripped.startswith("```"):
        block = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            block.append(lines[i]); i += 1
        add_code_block(block); i += 1; continue

    # table: current line looks like a row and next line is a --- separator
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i+1]) and "-" in lines[i+1]:
        header = split_table_row(lines[i]); i += 2
        rows = [header]
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_table_row(lines[i])); i += 1
        add_table(rows); doc.add_paragraph(); continue

    # horizontal rule
    if re.match(r"^-{3,}$", stripped):
        i += 1; continue

    # headings
    if stripped.startswith("#"):
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        level = len(m.group(1)); title = m.group(2)
        h = doc.add_heading(level=min(level, 4))
        add_runs(h, title)
        i += 1; continue

    # blockquote
    if stripped.startswith(">"):
        quote_lines = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            quote_lines.append(lines[i].strip()[1:].strip()); i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, " ".join(quote_lines).strip())
        for r in p.runs: r.italic = True
        continue

    # ordered list
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        p = doc.add_paragraph(style="List Number"); add_runs(p, m.group(2)); i += 1; continue
    # unordered list
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, stripped[2:]); i += 1; continue

    # blank
    if stripped == "":
        i += 1; continue

    # normal paragraph
    p = doc.add_paragraph(); add_runs(p, stripped); i += 1

doc.save(OUT)
print("Saved:", OUT)
